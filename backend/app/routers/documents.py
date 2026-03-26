from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.document import Document
from app.core.dependencies import get_current_user, require_leader_or_admin
from app.models.user import User
from app.models.audit import AuditLog
from app.services.openai_service import summarize_document
from app.services.rag_service import add_document_to_store, answer_question
import os
import fitz
import docx
from pydantic import BaseModel

router = APIRouter()
UPLOAD_DIR = "uploads"

def extract_text(file_path: str, ext: str) -> str:
    if ext == ".pdf":
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    elif ext == ".docx":
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    return ""

@router.post("/upload")
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    ALLOWED_EXTENSIONS = {'.pdf', '.txt', '.docx'}
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, "Only PDF, TXT, and DOCX files are allowed")

    MAX_SIZE = 10 * 1024 * 1024
    contents = await file.read()
    if len(contents) > MAX_SIZE:
        raise HTTPException(400, "File size must be under 10MB")
    await file.seek(0)
    
    import re
    safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename)
    file_path = f"uploads/{safe_filename}"
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    with open(file_path, "wb") as buffer:
        buffer.write(contents)
        
    ext = os.path.splitext(safe_filename)[1].lower()
    text = extract_text(file_path, ext)
    
    summary_data = summarize_document(text)
    
    new_doc = Document(
        filename=file.filename,
        file_path=file_path,
        raw_text=text,
        summary=summary_data.get("summary", ""),
        key_decisions=summary_data.get("key_decisions", ""),
        action_items=summary_data.get("action_items", ""),
        deadlines=summary_data.get("deadlines", ""),
        stakeholders=summary_data.get("stakeholders", ""),
        uploaded_by=int(current_user.get("sub", 0))
    )
    
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    add_document_to_store(str(new_doc.id), text, {"filename": safe_filename})
    
    audit = AuditLog(user_id=int(current_user.get("sub", 0)), user_name=current_user.get("name", "Unknown"), action="UPLOAD_DOCUMENT", module="Documents", details=f"Uploaded {safe_filename}")
    db.add(audit)
    db.commit()
    
    return {"id": new_doc.id, "summary": new_doc.summary, "message": "Uploaded successfully"}

@router.get("")
def list_documents(db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    return db.query(Document).order_by(Document.created_at.desc()).all()

@router.get("/{id}")
def get_document(id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@router.delete("/{id}")
def delete_document(id: int, db: Session = Depends(get_db), current_user: dict = Depends(require_leader_or_admin)):
    doc = db.query(Document).filter(Document.id == id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    
    db.delete(doc)
    
    audit = AuditLog(user_id=int(current_user.get("sub", 0)), user_name=current_user.get("name", "Unknown"), action="DELETE_DOCUMENT", module="Documents", details=f"Deleted doc {id}")
    db.add(audit)
    db.commit()
    return {"message": "Deleted"}

class QueryReq(BaseModel):
    question: str

@router.post("/query")
def query_docs(req: QueryReq, db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    ans = answer_question(req.question)
    audit = AuditLog(user_id=int(current_user.get("sub", 0)), user_name=current_user.get("name", "Unknown"), action="QUERY_DOCUMENTS", module="Documents", details=req.question)
    db.add(audit)
    db.commit()
    return {"answer": ans, "sources": []}
